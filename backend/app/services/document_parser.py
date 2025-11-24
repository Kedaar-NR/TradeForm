"""Service for parsing onboarding documents (PDF, DOCX, XLSX, CSV)."""

from pathlib import Path
from typing import Dict, Any, Optional
import logging
import json

logger = logging.getLogger(__name__)


class DocumentParserService:
    """Service for parsing various document types and extracting text/structured data."""
    
    @staticmethod
    def parse_document(file_path: str, mime_type: str) -> Dict[str, Any]:
        """
        Parse a document and extract text/structured data.
        
        Args:
            file_path: Path to the document file
            mime_type: MIME type of the document
            
        Returns:
            Dictionary with:
                - raw_text: Extracted text content
                - parsed_json: Optional structured data (for Excel/CSV)
                - success: Boolean indicating success
                - error: Optional error message
        """
        try:
            if mime_type == "application/pdf" or file_path.endswith('.pdf'):
                text = DocumentParserService.parse_pdf(file_path)
                return {
                    "raw_text": text,
                    "parsed_json": None,
                    "success": True,
                    "error": None
                }
            
            elif mime_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                              "application/msword"] or file_path.endswith(('.docx', '.doc')):
                text = DocumentParserService.parse_docx(file_path)
                return {
                    "raw_text": text,
                    "parsed_json": None,
                    "success": True,
                    "error": None
                }
            
            elif mime_type in ["application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                              "application/vnd.ms-excel"] or file_path.endswith(('.xlsx', '.xls')):
                result = DocumentParserService.parse_xlsx(file_path)
                return {
                    "raw_text": result["text"],
                    "parsed_json": result["json"],
                    "success": True,
                    "error": None
                }
            
            elif mime_type == "text/csv" or file_path.endswith('.csv'):
                result = DocumentParserService.parse_csv(file_path)
                return {
                    "raw_text": result["text"],
                    "parsed_json": result["json"],
                    "success": True,
                    "error": None
                }
            
            else:
                return {
                    "raw_text": "",
                    "parsed_json": None,
                    "success": False,
                    "error": f"Unsupported file type: {mime_type}"
                }
        
        except Exception as e:
            logger.error(f"Failed to parse document {file_path}: {str(e)}")
            return {
                "raw_text": "",
                "parsed_json": None,
                "success": False,
                "error": str(e)
            }
    
    @staticmethod
    def parse_pdf(file_path: str) -> str:
        """
        Extract text from PDF.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text as string
        """
        try:
            import pdfplumber
            
            text_parts = []
            
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text() or ""
                    text_parts.append(page_text)
            
            full_text = "\n\n".join(text_parts)
            logger.info(f"Successfully extracted {len(full_text)} characters from PDF")
            return full_text
        
        except ImportError:
            logger.error("pdfplumber not installed")
            raise Exception("PDF parsing library not available")
        except Exception as e:
            logger.error(f"Failed to parse PDF: {str(e)}")
            raise
    
    @staticmethod
    def parse_docx(file_path: str) -> str:
        """
        Extract text from DOCX.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text as string
        """
        try:
            from docx import Document
            
            doc = Document(file_path)
            
            # Extract text from paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            
            # Extract text from tables
            table_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        table_text.append(row_text)
            
            # Combine all text
            all_text = "\n\n".join(paragraphs)
            if table_text:
                all_text += "\n\n--- Tables ---\n\n" + "\n".join(table_text)
            
            logger.info(f"Successfully extracted {len(all_text)} characters from DOCX")
            return all_text
        
        except ImportError:
            logger.error("python-docx not installed")
            raise Exception("DOCX parsing library not available")
        except Exception as e:
            logger.error(f"Failed to parse DOCX: {str(e)}")
            raise
    
    @staticmethod
    def parse_xlsx(file_path: str) -> Dict[str, Any]:
        """
        Extract structured data from Excel.
        
        Args:
            file_path: Path to Excel file
            
        Returns:
            Dictionary with 'text' (formatted string) and 'json' (structured data)
        """
        try:
            import openpyxl
            
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            
            structured_data = {}
            text_parts = []
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                
                # Extract all rows
                rows = []
                for row in sheet.iter_rows(values_only=True):
                    # Filter out completely empty rows
                    if any(cell is not None for cell in row):
                        rows.append([str(cell) if cell is not None else "" for cell in row])
                
                if rows:
                    structured_data[sheet_name] = rows
                    
                    # Create text representation
                    text_parts.append(f"--- Sheet: {sheet_name} ---")
                    for row in rows[:100]:  # Limit to first 100 rows for text
                        text_parts.append(" | ".join(row))
            
            full_text = "\n".join(text_parts)
            json_data = json.dumps(structured_data)
            
            logger.info(f"Successfully parsed Excel with {len(structured_data)} sheets")
            return {
                "text": full_text,
                "json": json_data
            }
        
        except ImportError:
            logger.error("openpyxl not installed")
            raise Exception("Excel parsing library not available")
        except Exception as e:
            logger.error(f"Failed to parse Excel: {str(e)}")
            raise
    
    @staticmethod
    def parse_csv(file_path: str) -> Dict[str, Any]:
        """
        Extract structured data from CSV.
        
        Args:
            file_path: Path to CSV file
            
        Returns:
            Dictionary with 'text' (formatted string) and 'json' (structured data)
        """
        try:
            import csv
            
            rows = []
            
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            content = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding, newline='') as f:
                        reader = csv.reader(f)
                        rows = [row for row in reader if any(cell.strip() for cell in row)]
                    break
                except UnicodeDecodeError:
                    continue
            
            if not rows:
                raise Exception("Failed to read CSV with any encoding")
            
            # Create text representation
            text_parts = ["--- CSV Data ---"]
            for row in rows[:100]:  # Limit to first 100 rows
                text_parts.append(" | ".join(row))
            
            full_text = "\n".join(text_parts)
            json_data = json.dumps({"data": rows})
            
            logger.info(f"Successfully parsed CSV with {len(rows)} rows")
            return {
                "text": full_text,
                "json": json_data
            }
        
        except Exception as e:
            logger.error(f"Failed to parse CSV: {str(e)}")
            raise

