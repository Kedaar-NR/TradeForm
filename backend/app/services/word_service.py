"""Service for generating Word documents (.docx) from trade study reports."""

import io

from app.utils.text_parser import parse_report_blocks

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class WordService:
    """Service for generating Word documents from trade study reports."""

    @staticmethod
    def generate_report_docx(report_text: str) -> io.BytesIO:
        """
        Generate a formatted Word document from the trade study report text.
        
        Args:
            report_text: The markdown-formatted report text
            
        Returns:
            BytesIO object with .docx file
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is not installed. Install it with: pip install python-docx")

        # Parse report into structured blocks using shared utility
        blocks = parse_report_blocks(report_text)
        
        # Create document
        doc = Document()
        
        # Set document margins (1 inch all around)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Define color palette
        color_ink = RGBColor(15, 23, 42)      # #0f172a
        color_accent = RGBColor(31, 41, 55)   # #1f2937
        
        # Process blocks
        for block in blocks:
            block_type = block.get("type")
            
            if block_type == "heading":
                level = block.get("level", 1)
                text = str(block.get("text", ""))
                
                if level == 1:
                    heading = doc.add_heading(text, level=1)
                    for run in heading.runs:
                        run.font.size = Pt(18)
                        run.font.color.rgb = color_ink
                        run.font.bold = True
                    heading.space_after = Pt(10)
                else:
                    heading = doc.add_heading(text, level=2)
                    for run in heading.runs:
                        run.font.size = Pt(14)
                        run.font.color.rgb = color_accent
                        run.font.bold = True
                    heading.space_after = Pt(8)
                    
            elif block_type == "paragraph":
                text = str(block.get("text", ""))
                if text:
                    para = doc.add_paragraph(text)
                    for run in para.runs:
                        run.font.size = Pt(11)
                        run.font.name = 'Calibri'
                        run.font.color.rgb = color_ink
                    para.space_after = Pt(8)
                    
            elif block_type == "bullets":
                items = block.get("items", [])
                for item in items:
                    para = doc.add_paragraph(str(item), style='List Bullet')
                    for run in para.runs:
                        run.font.size = Pt(11)
                        run.font.name = 'Calibri'
                        run.font.color.rgb = color_ink
                    para.space_after = Pt(4)
        
        # Save to BytesIO buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer


def get_word_service() -> WordService:
    """Get WordService instance."""
    return WordService()
